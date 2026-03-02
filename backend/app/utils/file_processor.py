"""
Turli fayl turlaridan matn ajratib olish uchun fayl qayta ishlash vositalari.
"""
import os
from typing import Optional
from pathlib import Path
import io

# Ixtiyoriy importlar
try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    pytesseract = None
    Image = None

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

class FileProcessor:
    """Turli fayl turlaridan matn ajratib olish."""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.txt': 'txt',
        '.png': 'image',
        '.jpg': 'image',
        '.jpeg': 'image',
        '.gif': 'image',
        '.bmp': 'image',
        '.tiff': 'image',
    }
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """Fayl turi qo'llab-quvvatlanadimi tekshirish."""
        ext = Path(file_path).suffix.lower()
        return ext in FileProcessor.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def extract_text(file_path: str) -> Optional[str]:
        """
        Fayldan matn ajratib olish.
        
        Args:
            file_path: Fayl yo'li
        
        Returns:
            Ajratilgan matn yoki ajratish muvaffaqiyatsiz bo'lsa None
        """
        ext = Path(file_path).suffix.lower()
        file_type = FileProcessor.SUPPORTED_EXTENSIONS.get(ext)
        
        if not file_type:
            return None
        
        try:
            if file_type == 'pdf':
                return FileProcessor._extract_from_pdf(file_path)
            elif file_type == 'docx':
                return FileProcessor._extract_from_docx(file_path)
            elif file_type == 'txt':
                return FileProcessor._extract_from_txt(file_path)
            elif file_type == 'image':
                return FileProcessor._extract_from_image(file_path)
            else:
                return None
        except Exception:
            return None
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """PDF fayldan matn ajratib olish."""
        if not PYMUPDF_AVAILABLE:
            raise ValueError("PyMuPDF o'rnatilmagan. O'rnatish: pip install PyMuPDF")
        
        text_parts = []
        doc = fitz.open(file_path)
        
        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                text_parts.append(text)
                
                # Skaner qilingan PDF'lar uchun OCR'ni ham sinab ko'rish
                if not text.strip() and TESSERACT_AVAILABLE:
                    try:
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        ocr_text = pytesseract.image_to_string(img)
                        text_parts.append(ocr_text)
                    except Exception:
                        pass  # OCR muvaffaqiyatsiz, o'tkazib yuborish
        finally:
            doc.close()
        
        return '\n'.join(text_parts)
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """DOCX fayldan matn ajratib olish."""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx o'rnatilmagan. O'rnatish: pip install python-docx")
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # Jadvalardan ham ajratib olish
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """TXT fayldan matn ajratib olish."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def _extract_from_image(file_path: str) -> str:
        """OCR yordamida rasmdan matn ajratib olish."""
        if not TESSERACT_AVAILABLE:
            raise ValueError("Tesseract OCR o'rnatilmagan. O'rnatish: pip install pytesseract pillow")
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception:
            return ""

